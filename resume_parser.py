import os
import re
import logging
import zipfile
from pathlib import Path

logger = logging.getLogger(__name__)

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx 未安装，Word文档解析将不可用")

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("pdfplumber 未安装，PDF文档解析将不可用")

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logger.warning("pytesseract 或 Pillow 未安装，图片OCR解析将不可用")


class ResumeParser:
    def __init__(self):
        self.name_pattern = r'^([\u4e00-\u9fa5]{2,4})(?:老师|先生|女士)?$'
        self.phone_pattern = r'1[3-9]\d{9}'
        self.email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
        self.company_keywords = ['公司', '集团', '有限', '股份', '科技', '实业', '贸易', '投资', '企业']
        self.position_keywords = ['工程师', '经理', '总监', '主管', '专员', '顾问', '负责人', '总裁', '总经理', '董事长', 'CEO', 'CTO', 'COO', 'CFO']
        self.region_keywords = ['北京', '上海', '深圳', '广州', '杭州', '南京', '苏州', '成都', '武汉', '西安', '天津', '重庆', '省', '市', '区', '县']

    def extract_name(self, text):
        label_patterns = [
            r'(?:姓名|名字|Name)[:：\s]*([\u4e00-\u9fa5]{2,4})(?:老师|先生|女士)?',
            r'([\u4e00-\u9fa5]{2,4})(?:老师|先生|女士)?\s*(?:姓名|名字|Name)',
        ]

        for pattern in label_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1)

        lines = text.split('\n')
        for line in lines[:5]:
            line = line.strip()
            if not line:
                continue
            if len(line) >= 2 and len(line) <= 6:
                if re.match(r'^[\u4e00-\u9fa5]+$', line):
                    return line
            match = re.match(self.name_pattern, line)
            if match:
                return match.group(1)
        return None

    def extract_phone(self, text):
        matches = re.findall(self.phone_pattern, text)
        return list(set(matches))

    def extract_email(self, text):
        matches = re.findall(self.email_pattern, text)
        return list(set(matches))

    def extract_companies(self, text):
        companies = []
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            for keyword in self.company_keywords:
                if keyword in line:
                    match = re.search(r'([\u4e00-\u9fa5a-zA-Z0-9]+(?:公司|集团|有限|股份|科技|实业|贸易|投资)[\u4e00-\u9fa5a-zA-Z0-9]*)', line)
                    if match:
                        companies.append(match.group(1))
                    break
        return list(set(companies))

    def extract_position(self, text):
        positions = []
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            for keyword in self.position_keywords:
                if keyword in line:
                    positions.append(line)
                    break
        return list(set(positions))

    def extract_region(self, text):
        regions = []
        for keyword in self.region_keywords:
            if keyword in text:
                match = re.search(rf'([\u4e00-\u9fa5]+(?:省|市|区|县))', text)
                if match:
                    regions.append(match.group(1))
                break
        if not regions:
            for keyword in self.region_keywords:
                if len(keyword) >= 2:
                    match = re.search(rf'({keyword}[\u4e00-\u9fa5]*)', text)
                    if match:
                        regions.append(match.group(1))
                        break
        return list(set(regions))

    def parse_docx(self, filepath):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx 未安装，无法解析Word文档")
        
        text_parts = []
        doc = Document(filepath)
        for para in doc.paragraphs:
            text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        
        return '\n'.join(text_parts)

    def parse_pdf(self, filepath):
        if not PDF_AVAILABLE:
            raise ImportError("pdfplumber 未安装，无法解析PDF文档")
        
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        return '\n'.join(text_parts)

    def parse_image(self, filepath):
        if not OCR_AVAILABLE:
            raise ImportError("pytesseract 或 Pillow 未安装，无法解析图片")
        
        try:
            image = Image.open(filepath)
            text = pytesseract.image_to_string(image, lang='chi_sim')
            return text
        except Exception as e:
            logger.error(f"OCR解析图片失败: {e}")
            return ""

    def parse_txt(self, filepath):
        encodings = ['utf-8-sig', 'utf-8', 'gb18030', 'gbk', 'big5']
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _detect_file_kind(self, filepath):
        suffix = filepath.suffix.lower().lstrip('.')
        if suffix in {'docx', 'pdf', 'jpg', 'jpeg', 'png', 'bmp', 'txt'}:
            return suffix

        try:
            with open(filepath, 'rb') as f:
                header = f.read(16)
        except OSError:
            return suffix

        if header.startswith(b'PK\x03\x04'):
            return 'docx'
        if header.startswith(b'%PDF'):
            return 'pdf'
        if header.startswith(b'\xff\xd8\xff'):
            return 'jpg'
        if header.startswith(b'\x89PNG\r\n\x1a\n'):
            return 'png'
        if header.startswith(b'BM'):
            return 'bmp'

        return 'txt'

    def parse(self, filepath):
        filepath = Path(filepath)
        
        if not filepath.exists():
            raise FileNotFoundError(f"文件不存在: {filepath}")
        
        file_kind = self._detect_file_kind(filepath)
        
        if file_kind == 'docx':
            text = self.parse_docx(filepath)
        elif file_kind == 'pdf':
            text = self.parse_pdf(filepath)
        elif file_kind in ['jpg', 'jpeg', 'png', 'bmp']:
            text = self.parse_image(filepath)
        elif file_kind == 'txt':
            text = self.parse_txt(filepath)
        else:
            raise ValueError(f"不支持的文件格式: {filepath.suffix.lower()}")
        
        person_info = {
            'name': self.extract_name(text),
            'phone': self.extract_phone(text),
            'email': self.extract_email(text),
            'company': self.extract_companies(text),
            'position': self.extract_position(text),
            'region': self.extract_region(text),
            'raw_text': text
        }
        
        if person_info['phone']:
            person_info['phone'] = person_info['phone'][0] if len(person_info['phone']) == 1 else person_info['phone']
        if person_info['company']:
            person_info['company'] = person_info['company'][0] if len(person_info['company']) == 1 else person_info['company']
        if person_info['position']:
            person_info['position'] = person_info['position'][0] if len(person_info['position']) == 1 else person_info['position']
        if person_info['region']:
            person_info['region'] = person_info['region'][0] if len(person_info['region']) == 1 else person_info['region']
        
        return person_info


def parse_resume(filepath):
    parser = ResumeParser()
    return parser.parse(filepath)


if __name__ == '__main__':
    import sys
    if len(sys.argv) > 1:
        result = parse_resume(sys.argv[1])
        print("解析结果:")
        for key, value in result.items():
            if key != 'raw_text':
                print(f"  {key}: {value}")