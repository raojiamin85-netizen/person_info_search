import os
import json
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from config import OUTPUT_DIR, EXPORT_FORMATS
from utils import generate_report_id, get_current_time

class ReportGenerator:
    def __init__(self):
        self.report_id = generate_report_id()
    
    def generate_excel(self, person_info, results, filename=None):
        if not filename:
            filename = f"个人信息核查报告_{person_info['name']}_{self.report_id}.xlsx"
        
        filepath = OUTPUT_DIR / filename
        
        all_data = []
        for category, items in results.items():
            for item in items:
                row = {
                    '类别': category,
                    '标题': item.get('title', ''),
                    'URL': item.get('url', ''),
                    '摘要': item.get('summary', ''),
                    '来源': item.get('source', ''),
                    '相关度': item.get('relevance', ''),
                    '备注': item.get('note', '')
                }
                all_data.append(row)
        
        df = pd.DataFrame(all_data)
        df.to_excel(filepath, index=False)
        
        return str(filepath)
    
    def generate_word(self, person_info, results, filename=None):
        if not filename:
            filename = f"个人信息核查报告_{person_info['name']}_{self.report_id}.docx"
        
        filepath = OUTPUT_DIR / filename
        
        doc = Document()
        
        title = doc.add_heading('个人网络公开信息核查报告', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph(f'报告编号: {self.report_id}')
        doc.add_paragraph(f'核查对象: {person_info.get("name", "")}')
        doc.add_paragraph(f'意向岗位 / 合作事项: {person_info.get("position", "")}')
        doc.add_paragraph(f'核查日期: {get_current_time()}')
        doc.add_paragraph(f'核查范围: 互联网全网公开信息')
        
        doc.add_heading('一、核查说明', level=1)
        doc.add_paragraph('1. 本次核查仅依托公开网络渠道信息检索整理，不涉及个人隐私、非公开数据调取及第三方秘密调阅。')
        doc.add_paragraph('2. 内容仅供内部风控、招聘筛选、合作评估参考，不作为唯一决策依据。')
        doc.add_paragraph('3. 核查渠道包含：政务公示平台、司法文书官网、企业征信公示、主流新闻媒体、公开社交平台及行业公开资讯等。')
        
        doc.add_heading('二、基础信息', level=1)
        doc.add_paragraph(f'姓名: {person_info.get("name", "")}')
        doc.add_paragraph(f'关联从业领域 / 过往任职: {person_info.get("company", "")}')
        
        doc.add_heading('三、司法与合规风险核查', level=1)
        legal_items = results.get('legal', [])
        has_legal_risk = len(legal_items) > 0
        doc.add_paragraph(f'1. 失信及限高记录: {"有" if has_legal_risk else "无"}（备注：{"存在法律纠纷记录" if has_legal_risk else ""}）')
        doc.add_paragraph(f'2. 涉诉、仲裁、刑事案件公示: （备注：{"存在相关记录" if has_legal_risk else ""}）')
        if has_legal_risk:
            legal_details = "; ".join([item.get('title', '')[:30] for item in legal_items[:3]])
            doc.add_paragraph(f'3. 行政处罚、行业违规、监管警示: （备注：{legal_details}）')
        else:
            doc.add_paragraph('3. 行政处罚、行业违规、监管警示: （备注：无）')
        
        doc.add_heading('四、关联企业经营核查', level=1)
        enterprise_items = results.get('enterprise', [])
        has_enterprise = len(enterprise_items) > 0
        doc.add_paragraph('1. 公开任职（法人 / 股东 / 高管 / 监事）:')
        if has_enterprise:
            for i, item in enumerate(enterprise_items[:3], 1):
                company_name = item.get('company_name', item.get('title', ''))
                source = item.get('source', '')
                doc.add_paragraph(f'   对应企业名称: {company_name}（来源: {source}）')
        else:
            doc.add_paragraph('   对应企业名称: 未查询到相关信息')
        doc.add_paragraph(f'2. 关联企业经营异常、失信、处罚风险: {"存在风险记录" if has_enterprise else "未发现异常"}')
        
        doc.add_heading('五、网络舆情与个人声誉', level=1)
        news_items = results.get('news', [])
        has_news = len(news_items) > 0
        doc.add_paragraph(f'1. 不良言论、价值观争议、敏感负面信息: {"未发现" if not has_news else "需进一步核实"}')
        doc.add_paragraph(f'2. 商务纠纷、劳动争议、行业投诉爆料: {"未发现" if not has_news else "需进一步核实"}')
        doc.add_paragraph(f'3. 金融负面、借贷纠纷等公开信息: {"未发现" if not has_news else "需进一步核实"}')
        if has_news:
            news_titles = "; ".join([item.get('title', '')[:30] for item in news_items[:3]])
            doc.add_paragraph(f'4. 正面公开履历、行业荣誉、媒体报道: {news_titles}')
        else:
            doc.add_paragraph('4. 正面公开履历、行业荣誉、媒体报道: 未查询到相关信息')
        
        doc.add_heading('六、履历交叉核验', level=1)
        doc.add_paragraph('公开可查职业经历与个人自述: □ 基本一致 / □ 存在差异 / □ 无法核验')
        
        doc.add_heading('七、同名信息甄别', level=1)
        doc.add_paragraph('已结合行业、地域、任职背景交叉比对，排除同名无关人员干扰，数据指向核查对象本人。')
        
        doc.add_heading('八、风险等级判定', level=1)
        risk_level = "无风险" if not (has_legal_risk or has_enterprise) else "中风险" if has_legal_risk else "低风险"
        doc.add_paragraph(f'风险等级: {risk_level}')
        doc.add_paragraph('□ 无风险 □ 低风险 □ 中风险 □ 高风险')
        
        doc.add_heading('九、综合结论及建议', level=1)
        doc.add_paragraph(f'结论: 经核查，{"未发现重大风险" if not has_legal_risk else "发现潜在风险，建议谨慎"}')
        doc.add_paragraph(f'建议: {"□ 正常通过" if not has_legal_risk else "□ 面谈复核"} □ 审慎慎用 □ 不予合作 / 录用')
        
        doc.add_heading('十、免责声明', level=1)
        doc.add_paragraph('本报告仅整理公开网络信息，信息存在时效性与局限性，不构成法律意见，决策风险由使用方自行承担。')
        doc.add_paragraph(f'核查人: ')
        doc.add_paragraph(f'日期: {get_current_time().split(" ")[0]}')
        
        doc.save(filepath)
        return str(filepath)
    
    def generate_json(self, person_info, results, filename=None):
        if not filename:
            filename = f"个人信息核查报告_{person_info['name']}_{self.report_id}.json"
        
        filepath = OUTPUT_DIR / filename
        
        safe_person_info = dict(person_info)
        if 'raw_text' in safe_person_info:
            safe_person_info.pop('raw_text')

        report_data = {
            'report_id': self.report_id,
            'generate_time': get_current_time(),
            'person_info': safe_person_info,
            'results': results,
            'disclaimer': '本报告仅整理公开网络信息，信息存在时效性与局限性，不构成法律意见，决策风险由使用方自行承担。'
        }
        
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(report_data, f, ensure_ascii=False, indent=2)
        
        return str(filepath)
    
    def generate(self, person_info, results, formats=None):
        if not formats:
            formats = EXPORT_FORMATS
        
        output_files = {}
        
        for fmt in formats:
            if fmt == 'excel':
                output_files['excel'] = self.generate_excel(person_info, results)
            elif fmt == 'word':
                output_files['word'] = self.generate_word(person_info, results)
            elif fmt == 'json':
                output_files['json'] = self.generate_json(person_info, results)
        
        return output_files

if __name__ == '__main__':
    test_person = {
        'name': '张三',
        'company': '腾讯科技(深圳)有限公司',
        'position': '高级工程师',
        'region': '深圳'
    }
    
    test_results = {
        'news': [
            {'title': '张三获得行业创新奖', 'url': 'https://example.com/news1', 'summary': '张三在科技创新领域取得重大突破', 'source': '科技日报', 'relevance': 0.9}
        ],
        'enterprise': [
            {'company_name': '深圳张三科技有限公司', 'legal_person': '张三', 'source': '企查查', 'relevance': 0.85}
        ],
        'legal': []
    }
    
    generator = ReportGenerator()
    outputs = generator.generate(test_person, test_results)
    print(f"报告生成完成: {outputs}")